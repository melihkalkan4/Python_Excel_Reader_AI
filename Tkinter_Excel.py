from docx import Document
import pandas as pd
from tkinter import Tk, filedialog

def clean_percentage(value):
    try:
        return float(str(value).replace('%', '').replace(',', '.')) * 100
    except ValueError:
        return 0.0

def generate_text(df, filename, is_region=False):
    doc = Document()
    df = df.sort_values(['Bölge', 'Marka'])
    added_regions = set()
    added_marka = set()
    added_bricks = set()

    for index, row in df.iterrows():
        region = str(row['Bölge']) if pd.notna(row['Bölge']) else None
        marka = str(row['Marka']) if pd.notna(row['Marka']) else None
        brick = str(row['Brick']) if pd.notna(row['Brick']) else None

        if region and region != None and region not in added_regions:
            doc.add_heading(region if is_region else "Türkiye", level=1)
            added_regions.add(region)

        if marka and marka != None  and marka not in added_marka:
            doc.add_heading(f"Ürün Adı: {marka}", level=2)
            added_marka.add(marka)

        if brick and brick != None and brick not in added_bricks:
            doc.add_heading(brick, level=3)
            added_bricks.add(brick)
        
        text = f"""
        {marka} (YTD’DE) %{clean_percentage(row['PP_3.AY']):.2f} pazar payına sahiptir...
        """
        doc.add_paragraph(text)

    doc.save(filename)

def process_data(df_path, filename, is_region=False):
    df = pd.read_excel(df_path)
    generate_text(df, filename, is_region)

def select_file():
    root = Tk()
    root.withdraw()  # GUI penceresini gizle
    file_path = filedialog.askopenfilename()
    return file_path

def main():
    df_path = select_file()  # Kullanıcıdan dosya yolu seçmesini iste
    if df_path:
        process_data(df_path, "output.docx", is_region=False)

if __name__ == "__main__":
    main()
