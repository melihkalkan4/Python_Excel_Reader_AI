from docx import Document
import pandas as pd
import time

def clean_percentage(value):
    try:
        return float(str(value).replace('%', '').replace(',', '.')) * 100
    except ValueError:
        return 0.0  # Hatalı dönüşümde 0.0 dön

def generate_text(df, filename, is_region=False):
    start_time = time.time()
    doc = Document()
    df = df.sort_values(['Bölge', 'Marka'])
    print(f"Data sorted in {time.time() - start_time:.2f} seconds.")

    added_regions = set()
    added_marka = set()
    added_bricks = set()

    for index, row in df.iterrows():
        region = str(row['Bölge']) if pd.notna(row['Bölge']) else None
        marka = str(row['Marka']) if pd.notna(row['Marka']) else None
        brick = str(row['Brick']) if pd.notna(row['Brick']) else None

        if region and region != "Bilinmeyen Bölge" and region not in added_regions:
            doc.add_heading(region if is_region else "Türkiye", level=1)
            added_regions.add(region)
            print(f"Added heading for region: {region}")

        if marka and marka != "Bilinmeyen Marka" and marka not in added_marka:
            doc.add_heading(f"Ürün Adı: {marka}", level=2)
            added_marka.add(marka)
            print(f"Added heading for marka: {marka}")

        if brick and brick != "Bilinmeyen Brick" and brick not in added_bricks:
            doc.add_heading(brick, level=3)
            added_bricks.add(brick)
            print(f"Added heading for brick: {brick}")

        if region and marka and brick:
            text = f"""
            {marka} (YTD’DE) %{clean_percentage(row['PP_3.AY']):.2f} pazar payına sahiptir...
            """
            doc.add_paragraph(text)

    doc.save(filename)
    print(f"Document saved: {filename} in {time.time() - start_time:.2f} seconds.")

def process_data(df, filename, is_region=False):
    generate_text(df, filename, is_region)

def main():
    start_time = time.time()
    turkiye_df = pd.read_excel('C:\\Users\\melih.kalkan\\Desktop\\670B2000.xlsx', sheet_name='TURKIYE DATA')
    bolge_df = pd.read_excel('C:\\Users\\melih.kalkan\\Desktop\\670B2000.xlsx', sheet_name='BOLGE DATA')

    process_data(turkiye_df, "C:\\Users\\melih.kalkan\\Desktop\\Turkiye.docx")
    process_data(bolge_df, "C:\\Users\\melih.kalkan\\Desktop\\Bolge.docx", is_region=True)
    
    print(f"Total runtime: {time.time() - start_time:.2f} seconds.")

if __name__ == "__main__":
    main()

