import pandas as pd
from docx import Document

def clean_percentage(value):
    return float(str(value).replace('%', '').replace(',', '.')) * 100

def format_change(value):
    cleaned_value = clean_percentage(value)
    if cleaned_value < 0:
        return f"%{abs(cleaned_value):.2f} azalış"
    else:
        return f"%{cleaned_value:.2f} artış"

def generate_text(grouped_df, filename, is_region=False):
    doc = Document()
    for name, group in grouped_df:
        doc.add_heading(f"İlaç Adı: {name}", level=1)  # İlaç adına göre ana başlık oluşturuluyor

        for index, row in group.iterrows():
            if is_region == True:
                text = f"""{row['Bölge']}
                {row['Marka']}
                {row['Marka']} (YTD’DE) {format_change(row['PP_3.AY'])} pazar payına sahiptir. Geçtiğimiz üç aylık (Ocak-Şubat-Mart) dönemde pp değişimleri sırasıyla 
                {format_change(row['pp_değişim_1.ay'])}, 
                {format_change(row['pp_değişim_2.ay'])} ve 
                {format_change(row['pp_değişim_3.ay'])} olarak gerçekleşmiştir. 
                Bu süreçte {row['Marka']}'in Türkiye'deki ortalama artış oranı {format_change(row['pp_marg_ortalama'])} olarak kaydedilmiştir. 
                Geçtiğimiz yılın mart ayındaki Pazar payı yüzdesi {format_change(row['Mart_2023'])}'dır. Bu da önceki yılın aynı dönemine göre pazar payımızda {format_change(row['PP_3.AY']) - clean_percentage(row['Mart_2023'])} lik bir artışın sağlandığını göstermektedir.
                İncelemeye alınan son üç aylık dönemde Gelişim İndeksi sırasıyla 
                {format_change(row['GI_1.AY'])}, 
                {format_change(row['GI_2.AY'])} ve 
                {format_change(row['GI_3.AY'])} olarak tespit edilmiş ve ortalama olarak {format_change(row['GI_değişim_ortalama'])} hesaplanmıştır. 
                Mevcut trendlerin devam etmesi durumunda, 3 aylık süreçte {row['Marka']}'un net kutu hacminin {row['AylıkTahmin']} olacağı öngörülmektedir.
                """
            doc.add_paragraph(text)
        else:
            text = f"""{row['Marka']}
            {row['Marka']} (YTD’DE) {format_change(row['PP_3.AY'])} pazar payına sahiptir. Geçtiğimiz üç aylık (Ocak-Şubat-Mart) dönemde pp değişimleri sırasıyla 
            {format_change(row['pp_değişim_1.ay'])}, 
            {format_change(row['pp_değişim_2.ay'])} ve 
            {format_change(row['pp_değişim_3.ay'])} olarak gerçekleşmiştir. 
            Bu süreçte {row['Marka']}'in Türkiye'deki ortalama artış oranı {format_change(row['pp_marg_ortalama'])} olarak kaydedilmiştir. 
            Geçtiğimiz yılın mart ayındaki Pazar payı yüzdesi {format_change(row['Mart_2023'])}'dır. Bu da önceki yılın aynı dönemine göre pazar payımızda {format_change(row['PP_3.AY']) - clean_percentage(row['Mart_2023'])} lik bir artışın sağlandığını göstermektedir.
            İncelemeye alınan son üç aylık dönemde Gelişim İndeksi sırasıyla 
            {format_change(row['GI_1.AY'])}, 
            {format_change(row['GI_2.AY'])} ve 
            {format_change(row['GI_3.AY'])} olarak tespit edilmiş ve ortalama olarak {format_change(row['GI_değişim_ortalama'])} hesaplanmıştır. 
            Mevcut trendlerin devam etmesi durumunda, 3 aylık süreçte {row['Marka']}'un net kutu hacminin {row['AylıkTahmin']} olacağı öngörülmektedir.
            """
            doc.add_paragraph(text)
        doc.save(filename)

def process_data(df, group_column, filename, is_region=False):
    grouped_df = df.groupby(group_column)
    generate_text(grouped_df, filename, is_region)

def main(input_path_turkiye, input_path_bolge, output_path_turkiye, output_path_bolge):
    turkiye_df = pd.read_excel(input_path_turkiye, sheet_name='TURKIYE DATA')
    bolge_df = pd.read_excel(input_path_bolge, sheet_name='BOLGE DATA')

    process_data(turkiye_df, 'Marka', output_path_turkiye)
    process_data(bolge_df, 'Marka', output_path_bolge)

if __name__ == "__main__":
    main('C:\\Users\\melih.kalkan\\Desktop\\Turkiye_Bolge_Brick.xlsx', 
         'C:\\Users\\melih.kalkan\\Desktop\\Turkiye_Bolge_Brick.xlsx', 
         "C:\\Users\\melih.kalkan\\Desktop\\Turkiye.docx", 
         "C:\\Users\\melih.kalkan\\Desktop\\Bolge.docx")